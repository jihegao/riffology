#!/usr/bin/env python3
"""Build the screenshot-backed Riff user manual as a deterministic DOCX.

The Markdown source and five screenshots remain the editorial authority.  This
builder deliberately uses Word-native headings, numbering, tables, captions,
headers/footers, and image alternative text so the output remains editable and
auditable.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

# Songti.ttc is present on this macOS host and is made visible to the bundled
# headless LibreOffice renderer by the documented build/QA command. STSong is a
# Chinese-capable family within that collection; using the same family in both
# channels prevents mixed-script gaps in PDF export while retaining editable
# Word text. The preset's sizes and rhythm remain unchanged.
LATIN_FONT = "STSong"
EAST_ASIA_FONT = "STSong"
MONO_FONT = "STSong"

# Named visual override applied consistently to cover, headings, and quiet rules.
RIFF_TEAL_OVERRIDE = {
    "name": "riff_teal_accent",
    "teal": "0F766E",
    "teal_dark": "115E59",
    "teal_light": "E6F4F1",
    "ink": "163A3A",
    "muted": "5B6870",
    "line": "B8D9D4",
    "code_fill": "F3F7F6",
    "warning_fill": "FFF8E8",
}

TEAL = RGBColor.from_string(RIFF_TEAL_OVERRIDE["teal"])
TEAL_DARK = RGBColor.from_string(RIFF_TEAL_OVERRIDE["teal_dark"])
INK = RGBColor.from_string(RIFF_TEAL_OVERRIDE["ink"])
MUTED = RGBColor.from_string(RIFF_TEAL_OVERRIDE["muted"])
WHITE = RGBColor(255, 255, 255)


@dataclass(frozen=True)
class NumberingIds:
    bullet: int
    decimal: int
    checklist: int


def _set_font(run, *, name: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT,
              size: float | None = None, color: RGBColor | None = None,
              bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:hint"), "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, name: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT,
                    size: float | None = None, color: RGBColor | None = None,
                    bold: bool | None = None, italic: bool | None = None) -> None:
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:hint"), "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def _style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, TEAL_DARK),
        "Heading 2": (13, 14, 7, TEAL),
        "Heading 3": (12, 10, 5, INK),
    }
    for name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    # Custom cover styles are used instead of Word's built-in Title style.
    cover_title = _style(doc, "Manual Cover Title")
    _set_style_font(cover_title, size=29, color=INK, bold=True)
    cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(0)
    cover_title.paragraph_format.space_after = Pt(10)
    cover_title.paragraph_format.line_spacing = 1.0
    cover_title.paragraph_format.keep_together = True

    subtitle = _style(doc, "Manual Cover Subtitle")
    _set_style_font(subtitle, size=14, color=TEAL, bold=False)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle.paragraph_format.line_spacing = 1.1

    kicker = _style(doc, "Manual Kicker")
    _set_style_font(kicker, size=9.5, color=TEAL, bold=True)
    kicker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(18)

    meta = _style(doc, "Manual Cover Meta")
    _set_style_font(meta, size=10.5, color=MUTED)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(5)
    meta.paragraph_format.line_spacing = 1.15

    lead = _style(doc, "Manual Lead")
    _set_style_font(lead, size=11.5, color=INK, bold=False)
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(9)
    lead.paragraph_format.line_spacing = 1.3

    list_style = _style(doc, "Manual List")
    _set_style_font(list_style, size=11, color=INK)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25
    list_style.paragraph_format.widow_control = True

    code = _style(doc, "Manual Code")
    _set_style_font(code, name=MONO_FONT, east_asia=EAST_ASIA_FONT, size=8.8,
                    color=RGBColor.from_string("203434"))
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.08
    code.paragraph_format.keep_together = True

    callout = _style(doc, "Manual Callout")
    _set_style_font(callout, size=10.5, color=INK)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.22
    callout.paragraph_format.keep_together = True

    caption = _style(doc, "Manual Figure Caption")
    _set_style_font(caption, size=9, color=MUTED, italic=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.12
    caption.paragraph_format.keep_together = True

    table_text = _style(doc, "Manual Table Text")
    _set_style_font(table_text, size=9.2, color=INK)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.12

    table_header = _style(doc, "Manual Table Header")
    _set_style_font(table_header, size=9.2, color=TEAL_DARK, bold=True)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.12

    footer_style = _style(doc, "Manual Footer")
    _set_style_font(footer_style, size=8.5, color=MUTED)
    footer_style.paragraph_format.space_before = Pt(0)
    footer_style.paragraph_format.space_after = Pt(0)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def _set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_paragraph_border(paragraph, *, side: str, color: str, size: int = 10,
                          space: int = 4) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    edge = borders.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        borders.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    _set_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), RIFF_TEAL_OVERRIDE["muted"])
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    _set_font(run, size=8.5, color=MUTED)


def configure_headers_and_footers(doc: Document, validation_date: str) -> None:
    section = doc.sections[0]

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("RIFF  /  使用说明书")
    _set_font(left, size=8.5, color=TEAL, bold=True)
    right = p.add_run("    从空白 Project 到大样本仿真结论")
    _set_font(right, size=8.5, color=MUTED)

    first_header = section.first_page_header
    fp = first_header.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("RIFF  ·  实机使用说明书")
    _set_font(fr, size=8.5, color=TEAL, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.style = doc.styles["Manual Footer"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p)

    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0]
    fp.style = doc.styles["Manual Footer"]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"实机验收版  |  {validation_date}")
    _set_font(r, size=8.5, color=MUTED)


def _next_num_id(numbering) -> tuple[int, int]:
    root = numbering.element
    abs_ids = [int(el.get(qn("w:abstractNumId"))) for el in root.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in root.findall(qn("w:num"))]
    return (max(abs_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def _add_numbering_definition(doc: Document, *, fmt: str, text: str,
                              marker_dxa: int = 269, text_dxa: int = 540,
                              hanging_dxa: int = 271, font: str | None = None) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = _next_num_id(doc.part.numbering_part)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(text_dxa))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(text_dxa))
    ind.set(qn("w:hanging"), str(hanging_dxa))
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)

    if font:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
        rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        rpr.append(rfonts)
        lvl.append(rpr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def configure_numbering(doc: Document) -> NumberingIds:
    return NumberingIds(
        bullet=_add_numbering_definition(doc, fmt="bullet", text="•", font=EAST_ASIA_FONT),
        decimal=_add_numbering_definition(doc, fmt="decimal", text="%1.", font=EAST_ASIA_FONT),
        checklist=_add_numbering_definition(doc, fmt="bullet", text="✓", font=EAST_ASIA_FONT),
    )


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def _cell_margins(cell, margins=CELL_MARGIN_DXA) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side, value in margins.items():
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_fill(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def _fixed_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            _cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _column_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    ncols = len(rows[0])
    if ncols == 2:
        return [2700, 6660]
    if ncols == 3:
        return [3450, 1200, 4710]
    if ncols == 8:
        # Keep six-character extrema such as 0.9732 and 0.9944 on one line.
        # The P95 and KPI columns retain enough room for their longer labels.
        return [2000, 700, 1000, 1350, 850, 850, 1760, 850]

    scores = []
    for idx in range(ncols):
        values = [row[idx] if idx < len(row) else "" for row in rows]
        weighted = max(5.0, min(36.0, max(sum(2.0 if ord(ch) > 127 else 1.0 for ch in v) for v in values)))
        scores.append(weighted)
    minimum = 700
    remainder = CONTENT_WIDTH_DXA - minimum * ncols
    total = sum(scores)
    widths = [minimum + int(remainder * score / total) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(paragraph, text: str, *, base_size: float | None = None,
               base_color: RGBColor = INK, base_bold: bool | None = None) -> None:
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            _set_font(run, size=base_size, color=base_color, bold=base_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_font(run, name=MONO_FONT, east_asia=EAST_ASIA_FONT,
                      size=base_size or 9.3, color=TEAL_DARK)
            rpr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), RIFF_TEAL_OVERRIDE["teal_light"])
            rpr.append(shd)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_font(run, size=base_size, color=base_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_font(run, size=base_size, color=base_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_font(run, size=base_size, color=base_color, bold=base_bold)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    widths = _column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.style = doc.styles["Manual Table Header" if r_idx == 0 else "Manual Table Text"]
            if c_idx > 0 and len(value) < 28:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, base_size=9.2,
                       base_color=TEAL_DARK if r_idx == 0 else INK,
                       base_bold=True if r_idx == 0 else None)
            if r_idx == 0:
                _set_cell_fill(cell, RIFF_TEAL_OVERRIDE["teal_light"])
    _set_repeat_table_header(table.rows[0])
    _fixed_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr()
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.add_run("")


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    p = doc.add_paragraph(style="Manual Code")
    _set_paragraph_shading(p, RIFF_TEAL_OVERRIDE["code_fill"])
    _set_paragraph_border(p, side="left", color=RIFF_TEAL_OVERRIDE["teal"], size=18, space=5)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        _set_font(run, name=MONO_FONT, east_asia=EAST_ASIA_FONT, size=8.8,
                  color=RGBColor.from_string("203434"))


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Manual Callout")
    _set_paragraph_shading(p, RIFF_TEAL_OVERRIDE["warning_fill"])
    _set_paragraph_border(p, side="left", color="C89211", size=18, space=5)
    add_inline(p, text, base_size=10.5)


def _set_image_alt(inline_shape, alt: str) -> None:
    docpr = inline_shape._inline.docPr
    docpr.set("descr", alt)
    docpr.set("title", alt.split("：", 1)[0])


def add_figure(doc: Document, image_path: Path, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.2))
    _set_image_alt(shape, caption_text)
    caption = doc.add_paragraph(style="Manual Figure Caption")
    add_inline(caption, caption_text, base_size=9, base_color=MUTED)


def parse_table(lines: Sequence[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    if not lines[start].lstrip().startswith("|"):
        return None
    separator = lines[start + 1].strip()
    if not re.match(r"^\|?\s*:?-{3,}", separator):
        return None
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if i != start + 1:
            cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_cover(doc: Document, title: str, metadata: Sequence[str]) -> None:
    p = doc.add_paragraph(style="Manual Kicker")
    p.paragraph_format.space_before = Pt(82)
    p.add_run("RIFF  ·  实机闭环指南")

    p = doc.add_paragraph(style="Manual Cover Title")
    add_inline(p, title, base_size=29, base_color=INK, base_bold=True)

    p = doc.add_paragraph(style="Manual Cover Subtitle")
    add_inline(p, "从自然语言建模需求，到可视化运行、200 样本实验与持久结论",
               base_size=14, base_color=TEAL)

    rule = doc.add_paragraph()
    rule.paragraph_format.left_indent = Inches(1.1)
    rule.paragraph_format.right_indent = Inches(1.1)
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(26)
    _set_paragraph_border(rule, side="top", color=RIFF_TEAL_OVERRIDE["line"], size=8, space=1)

    for item in metadata:
        p = doc.add_paragraph(style="Manual Cover Meta")
        add_inline(p, item, base_size=10.5, base_color=MUTED)

    p = doc.add_paragraph(style="Manual Cover Meta")
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "空白 Project  →  需求落盘  →  可视化仿真  →  批量实验  →  证据化结论",
               base_size=9.5, base_color=TEAL, base_bold=True)
    doc.add_page_break()


def build_document(source: Path, output: Path) -> dict:
    text = source.read_text(encoding="utf-8")
    source_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    lines = text.splitlines()

    title = lines[0].removeprefix("# ").strip()
    metadata: list[str] = []
    idx = 1
    while idx < len(lines):
        stripped = lines[idx].strip()
        if stripped.startswith(">"):
            item = stripped[1:].strip()
            if item:
                metadata.append(item)
        elif stripped:
            break
        idx += 1
    validation_date = next((m.split("：", 1)[1] for m in metadata if m.startswith("验收日期：")), "2026-08-13")

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    numbering = configure_numbering(doc)
    configure_headers_and_footers(doc, validation_date)

    doc.core_properties.title = title
    doc.core_properties.subject = "从空白 Project 到自然语言建模、可视化仿真、大样本实验和分析结论的完整使用流程"
    doc.core_properties.author = "Riff"
    doc.core_properties.keywords = "Riff, Project, 仿真, 可视化, batch, 大样本, 使用说明书"
    doc.core_properties.comments = (
        "Preset=compact_reference_guide; Header=editorial_cover; "
        f"NamedOverride={RIFF_TEAL_OVERRIDE['name']}"
    )

    add_cover(doc, title, metadata)

    in_code = False
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []
    figure_count = 0
    decimal_list_active = False
    decimal_list_used = False
    active_decimal_num = numbering.decimal

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph_text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if not paragraph_text:
            return
        p = doc.add_paragraph(style="Manual Lead" if len(doc.paragraphs) < 10 else "Normal")
        add_inline(p, paragraph_text)

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            decimal_list_active = False
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        table = parse_table(lines, idx)
        if table:
            flush_paragraph()
            decimal_list_active = False
            rows, idx = table
            add_table(doc, rows)
            continue

        image_match = re.match(r"^!\[(.+)]\((.+)\)$", stripped)
        if image_match:
            flush_paragraph()
            decimal_list_active = False
            caption, rel_path = image_match.groups()
            image_path = (source.parent / rel_path).resolve()
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            add_figure(doc, image_path, caption)
            figure_count += 1
            idx += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            decimal_list_active = False
            level = min(3, len(heading_match.group(1)) - 1)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading_match.group(2),
                       base_size={1: 16, 2: 13, 3: 12}[level],
                       base_color={1: TEAL_DARK, 2: TEAL, 3: INK}[level],
                       base_bold=True)
            idx += 1
            continue

        checklist_match = re.match(r"^- \[[xX]\]\s+(.+)$", stripped)
        if checklist_match:
            flush_paragraph()
            decimal_list_active = False
            p = doc.add_paragraph(style="Manual List")
            apply_numbering(p, numbering.checklist)
            add_inline(p, checklist_match.group(1))
            idx += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            decimal_list_active = False
            p = doc.add_paragraph(style="Manual List")
            apply_numbering(p, numbering.bullet)
            add_inline(p, bullet_match.group(1))
            idx += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            flush_paragraph()
            if not decimal_list_active:
                if decimal_list_used:
                    active_decimal_num = _add_numbering_definition(
                        doc, fmt="decimal", text="%1."
                    )
                decimal_list_used = True
                decimal_list_active = True
            p = doc.add_paragraph(style="Manual List")
            apply_numbering(p, active_decimal_num)
            add_inline(p, ordered_match.group(1))
            idx += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            decimal_list_active = False
            quote_parts = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_parts.append(lines[idx].strip()[1:].strip())
                idx += 1
            add_callout(doc, " ".join(quote_parts))
            continue

        if stripped == "":
            flush_paragraph()
            decimal_list_active = False
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        decimal_list_active = False
        idx += 1

    flush_paragraph()

    # Remove direct use risk: no paragraph may use Word's built-in Title style.
    for p in doc.paragraphs:
        if p.style and p.style.name == "Title":
            raise RuntimeError("built-in Title style is prohibited")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    return {
        "source": str(source),
        "sourceSha256": source_hash,
        "output": str(output),
        "figureCount": figure_count,
        "tableCount": len(doc.tables),
        "paragraphCount": len(doc.paragraphs),
        "namedOverride": RIFF_TEAL_OVERRIDE,
    }


def audit_docx(path: Path, expected_figures: int = 6) -> dict:
    doc = Document(path)
    failures: list[str] = []
    section = doc.sections[0]

    def near(actual, expected, tolerance=4):
        return abs(actual - expected) <= tolerance

    if not near(section.page_width.twips, 12240):
        failures.append(f"page width {section.page_width.twips}")
    if not near(section.page_height.twips, 15840):
        failures.append(f"page height {section.page_height.twips}")
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        if not near(getattr(section, side).twips, 1440):
            failures.append(f"{side}={getattr(section, side).twips}")

    used_styles = {p.style.name for p in doc.paragraphs if p.style}
    if "Title" in used_styles:
        failures.append("built-in Title style is used")
    for required in ("Normal", "Heading 1", "Heading 2", "Heading 3",
                     "Manual Cover Title", "Manual Cover Subtitle",
                     "Manual Figure Caption"):
        if required not in {s.name for s in doc.styles}:
            failures.append(f"missing style {required}")

    with zipfile.ZipFile(path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        numbering_xml = package.read("word/numbering.xml").decode("utf-8")
        styles_xml = package.read("word/styles.xml").decode("utf-8")
        alt_count = len(re.findall(r'<wp:docPr[^>]+descr="[^\"]+"', document_xml))
        if alt_count != expected_figures:
            failures.append(f"image alt count {alt_count}, expected {expected_figures}")
        if numbering_xml.count("<w:abstractNum") < 3:
            failures.append("fewer than three numbering definitions")
        if f'w:eastAsia="{EAST_ASIA_FONT}"' not in styles_xml:
            failures.append("Chinese eastAsia font missing in styles")
        if "Manual Figure Caption" not in styles_xml:
            failures.append("caption style absent")

        table_count = document_xml.count("<w:tbl>")
        if table_count != len(doc.tables):
            failures.append("table XML count mismatch")
        for idx, table_xml in enumerate(re.findall(r"<w:tbl>.*?</w:tbl>", document_xml), start=1):
            if f'<w:tblW w:w="{CONTENT_WIDTH_DXA}" w:type="dxa"' not in table_xml and \
               f'<w:tblW w:type="dxa" w:w="{CONTENT_WIDTH_DXA}"' not in table_xml:
                failures.append(f"table {idx} missing 9360 DXA width")
            if f'<w:tblInd w:w="{TABLE_INDENT_DXA}" w:type="dxa"' not in table_xml and \
               f'<w:tblInd w:type="dxa" w:w="{TABLE_INDENT_DXA}"' not in table_xml:
                failures.append(f"table {idx} missing 120 DXA indent")
            if '<w:tblLayout w:type="fixed"' not in table_xml:
                failures.append(f"table {idx} layout not fixed")
            if '<w:tblHeader w:val="true"' not in table_xml:
                failures.append(f"table {idx} missing repeating header")

    report = {
        "path": str(path),
        "page": "US Letter portrait",
        "margins": "1in",
        "usedStyles": sorted(used_styles),
        "tableCount": len(doc.tables),
        "imageAltCount": expected_figures,
        "failures": failures,
        "passed": not failures,
    }
    if failures:
        raise RuntimeError(json.dumps(report, ensure_ascii=False, indent=2))
    return report


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--audit-json", type=Path)
    parser.add_argument("--expected-figures", type=int, default=6)
    args = parser.parse_args()

    build_report = build_document(args.source.resolve(), args.output.resolve())
    audit_report = audit_docx(args.output.resolve(), expected_figures=args.expected_figures)
    report = {"build": build_report, "audit": audit_report}
    if args.audit_json:
        args.audit_json.parent.mkdir(parents=True, exist_ok=True)
        args.audit_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
